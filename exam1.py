import streamlit as st
import pdfplumber
from docx import Document
from io import BytesIO
from openai import OpenAI

# ---------------------------
# Configuration
# ---------------------------
API_KEY = "AIzaSyByBDsl8gqHfwnzNXWYtkOzjCob3KROV3I"

client = OpenAI(
    api_key=API_KEY,
    base_url="https://generativelanguage.googleapis.com/v1beta/openai/"
)


# ---------------------------
# Helper functions
# ---------------------------
def load_notes_from_pdf(pdf_file):
    """Extract text from uploaded PDF using pdfplumber."""
    text = ""
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
    return text.strip()


def generate_exam(notes):
    """Call Gemini API to generate exam and marking guide."""
    exam_structure = """
    Set an exam with the following structure:
    - 20 multiple choice questions (1 mark each)
    - 40 short answer questions (5 marks each)
    - 30 essay question (10 marks)
    Provide a marking guide with expected answers and marks allocation.
    """

    response = client.chat.completions.create(
        model="gemini-2.5-flash",
        reasoning_effort="low",
        messages=[
            {"role": "system",
             "content": "You are an experienced lecturer who generates exam papers and detailed marking guides."},
            {"role": "user", "content": f"Lecture Notes:\n{notes}\n\n{exam_structure}"}
        ]
    )

    return response.choices[0].message.content


def save_to_word(text):
    """Save generated exam text to a Word file and return as BytesIO."""
    doc = Document()
    doc.add_heading("Exam Paper and Marking Guide", level=1)
    for line in text.split("\n"):
        doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------
# Streamlit UI
# ---------------------------
st.set_page_config(page_title="Exam Generator", page_icon="📘")

st.title("📘 AI Exam Paper Generator")

uploaded_file = st.file_uploader("Upload lecture notes (PDF)", type=["pdf"])

manual_notes = st.text_area("Or paste notes manually (if no PDF uploaded):")

if st.button("Generate Exam"):
    if uploaded_file:
        notes = load_notes_from_pdf(uploaded_file)
    elif manual_notes.strip():
        notes = manual_notes.strip()
    else:
        st.error("Please upload a PDF or paste notes manually.")
        st.stop()

    with st.spinner("Generating exam and marking guide..."):
        exam_text = generate_exam(notes)

    st.subheader("📄 Generated Exam and Marking Guide")
    st.text_area("Output", exam_text, height=400)

    # Save as Word
    word_file = save_to_word(exam_text)
    st.download_button(
        label="⬇️ Download as Word",
        data=word_file,
        file_name="exam_and_marking_guide.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )